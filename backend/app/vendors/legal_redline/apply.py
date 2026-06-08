"""Apply tracked changes (redlines) to Word documents via OOXML manipulation."""

import copy
import re
from datetime import datetime, timezone

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_rev_id_counter = 100000


def next_rev_id():
    global _rev_id_counter
    _rev_id_counter += 1
    return str(_rev_id_counter)


def enable_track_revisions(doc):
    """Set the trackRevisions flag in document settings."""
    settings = doc.settings.element
    if settings.find(qn("w:trackRevisions")) is None:
        track_rev = OxmlElement("w:trackRevisions")
        settings.append(track_rev)


def _make_ins(text, author, date_str, rpr=None):
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), next_rev_id())
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    ins.append(r)
    return ins


def _make_del(text, author, date_str, rpr=None):
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), next_rev_id())
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    r.append(dt)
    del_elem.append(r)
    return del_elem


def _make_trailing_run(text, rpr=None):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


def _get_full_paragraph_text(para):
    return "".join(run.text or "" for run in para.runs)


def _normalize_text(text):
    """Normalize text for fuzzy matching.

    Handles common mismatches between JSON input and document content:
    - Smart quotes vs straight quotes (common in legal docs and PDF conversions)
    - Multiple whitespace characters (tabs, double spaces from PDF line breaks)
    - En/em dashes vs hyphens
    """
    text = text.replace('\u2018', "'").replace('\u2019', "'")  # smart single quotes
    text = text.replace('\u201C', '"').replace('\u201D', '"')  # smart double quotes
    text = text.replace('\u2013', '-').replace('\u2014', '-')  # en/em dashes
    return re.sub(r'\s+', ' ', text)


def _find_text_across_runs(para, search_text):
    """
    Find search_text across potentially split runs.

    Uses a two-pass strategy:
    1. Exact substring match (fast path)
    2. Normalized match (handles smart quotes, extra whitespace from PDF conversion)

    When the normalized path matches, maps positions back to the original text
    so the correct character ranges are used for run splitting.

    Returns (start_run_idx, start_offset, end_run_idx, end_offset, matched_text)
    or None. matched_text is the actual document text that was matched (may differ
    from search_text in whitespace/quote characters).
    """
    runs = para.runs
    if not runs:
        return None

    char_map = []
    for i, run in enumerate(runs):
        for j in range(len(run.text or "")):
            char_map.append((i, j))

    if not char_map:
        return None

    full_text = _get_full_paragraph_text(para)
    matched_text = search_text

    # Try exact match first
    pos = full_text.find(search_text)

    # Fall back to normalized match (smart quotes, whitespace differences)
    if pos < 0:
        norm_full = _normalize_text(full_text)
        norm_search = _normalize_text(search_text)
        norm_pos = norm_full.find(norm_search)
        if norm_pos < 0:
            return None

        # Map normalized position back to original text position.
        # Walk both strings in parallel: the normalized string advances one char
        # at a time, while the original may have extra whitespace chars to skip.
        orig_idx = 0
        norm_idx = 0
        while norm_idx < norm_pos and orig_idx < len(full_text):
            if full_text[orig_idx].isspace():
                while orig_idx + 1 < len(full_text) and full_text[orig_idx + 1].isspace():
                    orig_idx += 1
            orig_idx += 1
            norm_idx += 1
        pos = orig_idx

        # Find the span length in original text that corresponds to norm_search
        search_len = 0
        search_norm_len = 0
        while search_norm_len < len(norm_search) and pos + search_len < len(full_text):
            search_len += 1
            if full_text[pos + search_len - 1].isspace():
                while pos + search_len < len(full_text) and full_text[pos + search_len].isspace():
                    search_len += 1
            search_norm_len += 1

        # Use the actual document text for the matched span
        matched_text = full_text[pos:pos + search_len]

    if pos < 0:
        return None

    start_run_idx, start_offset = char_map[pos]
    end_pos = pos + len(matched_text) - 1
    if end_pos >= len(char_map):
        return None
    end_run_idx, end_offset = char_map[end_pos]
    return (start_run_idx, start_offset, end_run_idx, end_offset, matched_text)


def _contains_normalized(full_text, search_text):
    """Check if search_text is in full_text, with whitespace-normalized fallback."""
    if search_text in full_text:
        return True
    return _normalize_text(search_text) in _normalize_text(full_text)


def apply_tracked_replacement(doc, search_text, replacement_text, author, date_str):
    """Replace search_text with replacement_text as tracked changes."""
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if not _contains_normalized(full_text, search_text):
            continue

        result = _find_text_across_runs(para, search_text)
        if result is None:
            continue

        start_run_idx, start_offset, end_run_idx, end_offset, matched_text = result
        runs = para.runs
        parent = para._element

        first_rpr = runs[start_run_idx]._element.find(qn("w:rPr"))

        end_run_text = runs[end_run_idx].text or ""
        trailing_text = end_run_text[end_offset + 1:]

        start_run_text = runs[start_run_idx].text or ""
        leading_text = start_run_text[:start_offset]

        all_children = list(parent)
        insert_pos = all_children.index(runs[start_run_idx]._element)

        for i in range(start_run_idx, end_run_idx + 1):
            parent.remove(runs[i]._element)

        new_elements = []
        if leading_text:
            new_elements.append(_make_trailing_run(leading_text, first_rpr))
        new_elements.append(_make_del(matched_text, author, date_str, first_rpr))
        new_elements.append(_make_ins(replacement_text, author, date_str, first_rpr))
        if trailing_text:
            new_elements.append(_make_trailing_run(trailing_text, first_rpr))

        for i, elem in enumerate(new_elements):
            parent.insert(insert_pos + i, elem)

        return True
    return False


def apply_tracked_deletion(doc, delete_text, author, date_str):
    """Delete text as a tracked deletion."""
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if not _contains_normalized(full_text, delete_text):
            continue

        result = _find_text_across_runs(para, delete_text)
        if result is None:
            continue

        start_run_idx, start_offset, end_run_idx, end_offset, matched_text = result
        runs = para.runs
        parent = para._element

        first_rpr = runs[start_run_idx]._element.find(qn("w:rPr"))

        end_run_text = (runs[end_run_idx].text or "")
        trailing_text = end_run_text[end_offset + 1:]

        start_run_text = (runs[start_run_idx].text or "")
        leading_text = start_run_text[:start_offset]

        insert_pos = list(parent).index(runs[start_run_idx]._element)

        for i in range(start_run_idx, end_run_idx + 1):
            parent.remove(runs[i]._element)

        new_elements = []
        if leading_text:
            new_elements.append(_make_trailing_run(leading_text, first_rpr))
        new_elements.append(_make_del(matched_text, author, date_str, first_rpr))
        if trailing_text:
            new_elements.append(_make_trailing_run(trailing_text, first_rpr))

        for i, elem in enumerate(new_elements):
            parent.insert(insert_pos + i, elem)

        return True
    return False


def apply_tracked_insertion(doc, after_text, new_text, author, date_str,
                            as_paragraph=False):
    """Insert new text after anchor text as a tracked insertion.

    Args:
        as_paragraph: If True, insert as a new paragraph after the anchor
            paragraph instead of inline. Use for whole clauses/sections.
    """
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if not _contains_normalized(full_text, after_text):
            continue

        result = _find_text_across_runs(para, after_text)
        if result is None:
            continue

        if as_paragraph:
            # Insert as a new paragraph after the anchor paragraph
            body = para._element.getparent()
            target_idx = list(body).index(para._element)

            new_para = OxmlElement("w:p")
            target_pPr = para._element.find(qn("w:pPr"))
            if target_pPr is not None:
                new_pPr = copy.deepcopy(target_pPr)
                for numPr in new_pPr.findall(qn("w:numPr")):
                    new_pPr.remove(numPr)
                new_para.append(new_pPr)

            end_rpr = None
            target_runs = para._element.findall(qn("w:r"))
            if target_runs:
                end_rpr = target_runs[0].find(qn("w:rPr"))

            ins_elem = _make_ins(new_text, author, date_str, end_rpr)
            new_para.append(ins_elem)
            body.insert(target_idx + 1, new_para)
            return True

        _, _, end_run_idx, end_offset, _ = result
        runs = para.runs
        parent = para._element

        end_run = runs[end_run_idx]
        end_rpr = end_run._element.find(qn("w:rPr"))
        end_run_text = end_run.text or ""

        trailing_text = end_run_text[end_offset + 1:]
        if trailing_text:
            end_run.text = end_run_text[:end_offset + 1]

        insert_pos = list(parent).index(end_run._element) + 1

        ins_elem = _make_ins(new_text, author, date_str, end_rpr)
        parent.insert(insert_pos, ins_elem)

        if trailing_text:
            parent.insert(insert_pos + 1, _make_trailing_run(trailing_text, end_rpr))

        return True
    return False


def apply_tracked_add_section(doc, after_section, new_text, author, date_str,
                               new_section_number=None):
    """Insert a new section (as tracked insertion) after a section reference."""
    # Find the paragraph containing the after_section reference
    target_para = None
    for para in doc.paragraphs:
        full_text = _get_full_paragraph_text(para)
        if _contains_normalized(full_text, after_section):
            target_para = para

    if target_para is None:
        return False

    # Build the new paragraph text
    if new_section_number:
        insert_text = f"{new_section_number}. {new_text}"
    else:
        insert_text = new_text

    # Create a new paragraph after the target as a tracked insertion
    parent = target_para._element.getparent()
    target_idx = list(parent).index(target_para._element)

    # Create new paragraph element with formatting from target paragraph
    new_para = OxmlElement("w:p")
    target_pPr = target_para._element.find(qn("w:pPr"))
    if target_pPr is not None:
        import copy as _copy
        new_pPr = _copy.deepcopy(target_pPr)
        # Remove numbering properties to avoid wrong auto-numbering
        for numPr in new_pPr.findall(qn("w:numPr")):
            new_pPr.remove(numPr)
        new_para.append(new_pPr)

    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), next_rev_id())
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), date_str)
    r = OxmlElement("w:r")
    # Copy run properties (font, size) from target paragraph's first run
    target_runs = target_para._element.findall(qn("w:r"))
    if target_runs:
        target_rPr = target_runs[0].find(qn("w:rPr"))
        if target_rPr is not None:
            import copy as _copy
            r.append(_copy.deepcopy(target_rPr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = insert_text
    r.append(t)
    ins.append(r)
    new_para.append(ins)

    parent.insert(target_idx + 1, new_para)
    return True


def apply_redlines(input_path, output_path, redlines, author="Reviewer"):
    """
    Apply a list of redlines to a .docx file.

    Args:
        input_path: Path to original .docx
        output_path: Path for output .docx with tracked changes
        redlines: List of dicts with type, old/new/text/anchor fields
        author: Author name for tracked changes

    Returns:
        List of (status, description) tuples
    """
    doc = Document(input_path)
    enable_track_revisions(doc)

    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    results = []
    for i, redline in enumerate(redlines):
        rtype = redline.get("type")
        if rtype is None:
            results.append(("ERROR", f"Redline {i}: missing 'type' field"))
            print(f"  [ERROR] Redline {i}: missing 'type' field")
            continue

        try:
            success = False

            if rtype == "replace":
                old = redline["old"]
                new = redline["new"]
                success = apply_tracked_replacement(doc, old, new, author, date_str)
                desc = f"Replace: '{old[:50]}...' -> '{new[:50]}...'"

            elif rtype == "delete":
                text = redline["text"]
                success = apply_tracked_deletion(doc, text, author, date_str)
                desc = f"Delete: '{text[:60]}...'"

            elif rtype == "insert_after":
                anchor = redline["anchor"]
                text = redline["text"]
                as_para = redline.get("as_paragraph", False)
                success = apply_tracked_insertion(
                    doc, anchor, text, author, date_str,
                    as_paragraph=as_para,
                )
                desc = f"Insert after: '{anchor[:40]}...'"

            elif rtype == "add_section":
                after_sec = redline.get("after_section", "")
                text = redline["text"]
                new_sec = redline.get("new_section_number")
                success = apply_tracked_add_section(
                    doc, after_sec, text, author, date_str, new_sec
                )
                desc = f"Add section: '{new_sec or 'new'}' after '{after_sec}'"

            else:
                desc = f"Unknown type: {rtype}"

            status = "OK" if success else "NOT FOUND"
            results.append((status, desc))
            print(f"  [{status}] {desc}")

        except KeyError as e:
            desc = f"Redline {i} ({rtype}): missing required field {e}"
            results.append(("ERROR", desc))
            print(f"  [ERROR] {desc}")

    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    print(f"Applied: {sum(1 for s, _ in results if s == 'OK')}/{len(results)} redlines")
    return results
