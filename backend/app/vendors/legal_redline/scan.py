"""Scan Word documents for blank fields, placeholders, and missing references."""

import re
from docx import Document


# Patterns that indicate blank/placeholder fields
PLACEHOLDER_PATTERNS = [
    (r'\$\s*[_X]+', "Dollar amount placeholder"),
    (r'\[[\s_]*\]', "Empty brackets"),
    (r'\[enter[^\]]*\]', "Enter-value placeholder"),
    (r'\[amount[^\]]*\]', "Amount placeholder"),
    (r'\[date[^\]]*\]', "Date placeholder"),
    (r'\[name[^\]]*\]', "Name placeholder"),
    (r'\[TBD\]|TBD', "To be determined"),
    (r'_{3,}', "Blank line"),
    (r'\bXX+\b', "XX placeholder"),
    (r'\$\s*XX', "Dollar XX placeholder"),
    (r'\[\s*%\s*\]|\[__%\]', "Percentage placeholder"),
    (r'EUR\s*0[,.]?0*\b', "Zero EUR amount"),
    (r'\bN/A\b(?=.*(?:fee|amount|rate|price))', "N/A in financial field"),
]

# Patterns for exhibit/schedule references
REFERENCE_PATTERN = re.compile(
    r'(?:Exhibit|Schedule|Annex|Appendix|Attachment)\s+([A-Z0-9]+(?:\.\d+)?)',
    re.IGNORECASE
)


def scan_placeholders(docx_path):
    """
    Scan a .docx for blank fields and placeholders.

    Returns list of dicts:
        {"paragraph": int, "text": str, "match": str, "pattern_desc": str}
    """
    doc = Document(docx_path)
    findings = []

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        if not text.strip():
            continue

        for pattern, desc in PLACEHOLDER_PATTERNS:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                context_start = max(0, match.start() - 30)
                context_end = min(len(text), match.end() + 30)
                context = text[context_start:context_end]
                if context_start > 0:
                    context = "..." + context
                if context_end < len(text):
                    context = context + "..."

                findings.append({
                    "paragraph": para_idx,
                    "match": match.group(),
                    "context": context,
                    "pattern_desc": desc,
                })

    # Also scan tables
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text or ""
                if not text.strip():
                    continue
                for pattern, desc in PLACEHOLDER_PATTERNS:
                    for match in re.finditer(pattern, text, re.IGNORECASE):
                        findings.append({
                            "table": tbl_idx,
                            "row": row_idx,
                            "cell": cell_idx,
                            "match": match.group(),
                            "context": text[:80],
                            "pattern_desc": desc,
                        })

    return findings


def scan_references(docx_path):
    """
    Scan a .docx for exhibit/schedule/appendix references and check which exist.

    Returns dict:
        {"referenced": set of names, "defined": set of names, "missing": set of names}
    """
    doc = Document(docx_path)

    referenced = set()
    defined = set()

    for para in doc.paragraphs:
        text = para.text or ""
        # Find references
        for match in REFERENCE_PATTERN.finditer(text):
            referenced.add(match.group(1).upper())

        # Check if paragraph defines an exhibit (e.g., "EXHIBIT A" as a heading)
        stripped = text.strip().upper()
        for prefix in ("EXHIBIT ", "SCHEDULE ", "ANNEX ", "APPENDIX ", "ATTACHMENT "):
            if stripped.startswith(prefix) and len(stripped) < len(prefix) + 10:
                name = stripped[len(prefix):].strip().rstrip(":")
                if name:
                    defined.add(name)

    return {
        "referenced": sorted(referenced),
        "defined": sorted(defined),
        "missing": sorted(referenced - defined),
    }


def scan_document(docx_path):
    """
    Run all scans on a document. Returns a full report dict.
    """
    placeholders = scan_placeholders(docx_path)
    refs = scan_references(docx_path)

    return {
        "placeholders": placeholders,
        "references": refs,
        "placeholder_count": len(placeholders),
        "missing_exhibit_count": len(refs["missing"]),
    }
